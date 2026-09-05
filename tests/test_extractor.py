"""extractor 与 CLI 端到端测试（用 python-docx 现场生成样例文档）。"""

from types import SimpleNamespace

from thesislint.cli import analyze
from thesislint.extractor import find_bibliography


def fake_paragraphs(*pairs):
    """pairs: (text, is_heading) -> 带样式属性的对象列表。"""
    return [
        SimpleNamespace(text=t, style=SimpleNamespace(name="Heading 1" if h else "Normal"))
        for t, h in pairs
    ]


class TestExtractor:
    def test_finds_section_between_headings(self):
        paras = fake_paragraphs(
            ("绪论", True),
            ("正文正文。", False),
            ("参考文献", True),
            ("[1] 张三. 文一[J]. 学报, 2024, 1(1): 1-9.", False),
            ("[2] 李四. 文二[M]. 北京: 出版社, 2020: 1-9.", False),
            ("致谢", True),
            ("感谢大家。", False),
        )
        bib = find_bibliography(paras)
        assert bib.found
        assert len(bib.entries) == 2
        assert bib.entries[0].startswith("[1] 张三")
        assert bib.line_numbers == [4, 5]

    def test_merges_wrapped_entry(self):
        paras = fake_paragraphs(
            ("参考文献", True),
            ("[3] 王五, 赵六. 一个标题特别长导致在 Word 里折行的条目[J]. 学报,", False),
            ("2024, 47(3): 100-110.", False),
        )
        bib = find_bibliography(paras)
        assert len(bib.entries) == 1
        assert "2024" in bib.entries[0]

    def test_merges_manual_line_break_inside_paragraph(self):
        paras = fake_paragraphs(
            ("参考文献", True),
            ("[1] 张三. 论文[J].\n学报, 2024, 1(1): 1-9.", False),
        )
        bib = find_bibliography(paras)
        assert bib.entries == ["[1] 张三. 论文[J]. 学报, 2024, 1(1): 1-9."]

    def test_numbered_heading_with_colon(self):
        paras = fake_paragraphs(
            ("第 5 章 参考文献：", False),
            ("[1] 张三. 论文[J]. 学报, 2024, 1(1): 1-9.", False),
        )
        bib = find_bibliography(paras)
        assert bib.found
        assert bib.heading_text == "第 5 章 参考文献："

    def test_deeper_heading_inside_section_is_skipped(self):
        paras = [
            SimpleNamespace(text="参考文献", style=SimpleNamespace(name="Heading 1")),
            SimpleNamespace(text="中文文献", style=SimpleNamespace(name="Heading 2")),
            SimpleNamespace(
                text="[1] 张三. 论文[J]. 学报, 2024, 1(1): 1-9.",
                style=SimpleNamespace(name="Normal"),
            ),
            SimpleNamespace(text="致谢", style=SimpleNamespace(name="Heading 1")),
        ]
        bib = find_bibliography(paras)
        assert len(bib.entries) == 1

    def test_accepts_paragraph_iterables(self):
        paras = iter(
            fake_paragraphs(
                ("参考文献", True),
                ("[1] 张三. 论文[J]. 学报, 2024, 1(1): 1-9.", False),
            )
        )
        assert len(find_bibliography(paras).entries) == 1

    def test_orphan_line_becomes_missing_number_entry(self):
        """区内开头就是无序号正文：必须成为条目交给 checker 报 E001，不得静默跳过。"""
        paras = fake_paragraphs(
            ("参考文献", True),
            ("张三. 丢了序号的条目[J]. 学报, 2024, 1(1): 1-9.", False),
            ("[2] 李四. 正常条目[J]. 学报, 2024, 2(1): 1-9.", False),
        )
        bib = find_bibliography(paras)
        assert len(bib.entries) == 2
        assert bib.entries[0].startswith("张三")

    def test_not_found(self):
        paras = fake_paragraphs(("只有正文", False))
        assert not find_bibliography(paras).found

    def test_empty_section_keyword_is_not_a_heading(self):
        paras = fake_paragraphs(("参考文献", True))
        assert not find_bibliography(paras, "").found

    def test_multiple_entries_in_one_paragraph(self):
        paras = fake_paragraphs(
            ("参考文献", True),
            ("[1] 张三. 文一[J].\n学报, 2024.\n[2] 李四. 文二[J]. 学报, 2023.", False),
            ("致谢", True),
        )
        bib = find_bibliography(paras)
        assert bib.entries == [
            "[1] 张三. 文一[J]. 学报, 2024.",
            "[2] 李四. 文二[J]. 学报, 2023.",
        ]
        assert bib.line_numbers == [2, 2]

    def test_inline_brackets_are_not_new_entries(self):
        entry = "[1] 张三. 对照 [2] 的方法[J]. 学报, 2024."
        bib = find_bibliography(fake_paragraphs(("参考文献", True), (entry, False)))
        assert bib.entries == [entry]


class TestCliAnalyze:
    def test_duplicate_number_after_manual_break_is_reported(self, tmp_path):
        from docx import Document

        doc = Document()
        doc.add_heading("参考文献", level=1)
        doc.add_paragraph("[1] 张三. 文一[J]. 学报, 2024.\n[1] 李四. 文二[J]. 学报, 2023.")
        path = tmp_path / "multiple-entries.docx"
        doc.save(path)
        report = analyze(path)
        assert len(report.entries) == 2
        assert any(issue.code == "S002" for issue in report.section_issues)

    def test_analyze_real_docx(self, tmp_path):
        from docx import Document

        doc = Document()
        doc.add_heading("参考文献", level=1)
        doc.add_paragraph("[1] 张三. 文一[J]. 学报, 2024, 1(1): 1-9.")
        doc.add_paragraph("[2] 李四. 文二. 学报, 2024.")  # 缺类型标识
        doc.add_paragraph("[2] 王五. 文三[J]. 学报, 2024, 1(1): 1-9.")  # 序号重复
        path = tmp_path / "sample.docx"
        doc.save(str(path))

        report = analyze(path)
        assert report.found_section
        assert len(report.entries) == 3
        assert any(i.code == "E002" for i in report.entries[1].issues)
        assert any(i.code == "S002" for i in report.section_issues)

    def test_clean_docx_zero_issues(self, tmp_path):
        from docx import Document

        doc = Document()
        doc.add_heading("参考文献", level=1)
        doc.add_paragraph("[1] 张三. 文一[J]. 学报, 2024, 1(1): 1-9.")
        path = tmp_path / "clean.docx"
        doc.save(str(path))

        report = analyze(path)
        assert report.error_count == 0
        assert report.warn_count == 0

    def test_real_docx_manual_line_break_is_fully_checked(self, tmp_path):
        from docx import Document

        doc = Document()
        doc.add_heading("参考文献", level=1)
        paragraph = doc.add_paragraph("[1] 张三. 文一[J].")
        paragraph.add_run().add_break()
        paragraph.add_run("学报, 2024, 1(1): 1-9.")
        path = tmp_path / "line-break.docx"
        doc.save(str(path))

        report = analyze(path)
        assert report.error_count == 0
        assert report.warn_count == 0
