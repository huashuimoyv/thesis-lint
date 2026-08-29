"""CLI 退出码契约测试。

背景：README 承诺 `--strict` 可用于 CI（有警告也算失败），
但 v0.4.0 中退出码路径（cli.py:87-101）零测试覆盖。
本文件把「退出码契约」钉死。
"""

from __future__ import annotations

import pytest
from docx import Document

from thesislint import gui
from thesislint.cli import main


@pytest.fixture(autouse=True)
def _force_cli_mode(monkeypatch):
    """强制走命令行分支，避免在装有 tkinter 的机器上误开 GUI 导致测试卡死。

    cli.main 内部是 `from .gui import should_launch_gui`，
    should_launch_gui 又调用模块级 _supports_tk()，所以 patch 模块属性即可生效。
    与 tests/test_modes.py 的做法保持一致。
    """
    monkeypatch.setattr(gui, "_supports_tk", lambda: False)


def _make_docx(path, entries: list[str], heading: str = "参考文献"):
    doc = Document()
    doc.add_heading(heading, level=1)
    for e in entries:
        doc.add_paragraph(e)
    doc.save(str(path))
    return path


class TestExitCodes:
    def test_clean_doc_returns_zero(self, tmp_path):
        path = _make_docx(
            tmp_path / "clean.docx",
            [
                "[1] 张三. 文一[J]. 学报, 2024, 1(1): 1-9.",
                "[2] 李四. 文二[M]. 北京: 出版社, 2020: 1-9.",
            ],
        )
        assert main([str(path)]) == 0

    def test_error_returns_one(self, tmp_path):
        path = _make_docx(
            tmp_path / "err.docx",
            [
                "张三. 无序号无类型的条目. 学报, 2024.",  # E001 + E002
            ],
        )
        assert main([str(path)]) == 1

    def test_warning_returns_zero_without_strict(self, tmp_path):
        path = _make_docx(
            tmp_path / "warn.docx",
            [
                "[1] 张三，李四. 论文名称[J]. 学报，2024，1(1): 1-9.",  # W005 全角
            ],
        )
        assert main([str(path)]) == 0

    def test_warning_returns_one_with_strict(self, tmp_path):
        path = _make_docx(
            tmp_path / "warn.docx",
            [
                "[1] 张三，李四. 论文名称[J]. 学报，2024，1(1): 1-9.",  # W005 全角
            ],
        )
        assert main([str(path), "--strict"]) == 1

    def test_duplicate_number_returns_one(self, tmp_path):
        path = _make_docx(
            tmp_path / "dup.docx",
            [
                "[1] 张三. 文一[J]. 学报, 2024, 1(1): 1-9.",
                "[1] 李四. 文二[J]. 学报, 2024, 2(1): 1-9.",  # S002
            ],
        )
        assert main([str(path)]) == 1


class TestUsageErrors:
    def test_missing_file_returns_two(self):
        assert main(["nonexistent.docx"]) == 2

    def test_wrong_extension_returns_two(self, tmp_path):
        p = tmp_path / "a.pdf"
        p.write_bytes(b"%PDF-1.4")
        assert main([str(p)]) == 2

    def test_no_section_returns_zero(self, tmp_path):
        path = _make_docx(tmp_path / "nosec.docx", ["只有正文"], heading="绪论")
        assert main([str(path)]) == 0
