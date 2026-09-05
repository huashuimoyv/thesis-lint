"""桌面工作流：文件传递、报告快照、过期回调与导出失败。"""

import sys
from pathlib import Path
from types import SimpleNamespace

import pytest
from docx import Document

from thesislint import cli, gui


def test_make_root_uses_supported_drag_and_drop_api(monkeypatch):
    registered = []
    root = SimpleNamespace(
        drop_target_register=lambda value: registered.append(value),
        dnd_bind=lambda event, callback: registered.append((event, callback)),
    )
    monkeypatch.setitem(
        sys.modules,
        "tkinterdnd2",
        SimpleNamespace(
            DND_FILES="DND_Files",
            TkinterDnD=SimpleNamespace(Tk=lambda: root),
        ),
    )
    app = gui.App.__new__(gui.App)
    fallback = object()
    assert app._make_root(SimpleNamespace(Tk=lambda: fallback)) is root
    assert registered == ["DND_Files", ("<<Drop>>", app._on_drop)]


def test_make_root_still_works_when_drag_library_is_unavailable(monkeypatch):
    monkeypatch.setitem(sys.modules, "tkinterdnd2", None)
    root = object()
    app = gui.App.__new__(gui.App)
    assert app._make_root(SimpleNamespace(Tk=lambda: root)) is root


def test_empty_document_is_not_shown_as_passed(tmp_path):
    path = tmp_path / "empty.docx"
    doc = Document()
    doc.add_heading("参考文献", level=1)
    doc.save(path)
    app = gui.App.__new__(gui.App)
    callbacks = []
    finished = []
    app.root = SimpleNamespace(after=lambda delay, callback: callbacks.append(callback))
    app._finish = lambda *args: finished.append(args)
    app._work(str(path), path.name, 1)
    callbacks.pop(0)()
    assert "未提取到" in finished[0][5]
    assert finished[0][6] != gui.OK


@pytest.mark.parametrize("args", [[], ["含空格的 论文.docx"]])
def test_explorer_launch_forwards_file(monkeypatch, args):
    received = []
    monkeypatch.setattr(gui, "should_launch_gui", lambda _args: True)
    monkeypatch.setattr(
        gui, "run_gui", lambda initial_path=None: received.append(initial_path) or 0
    )
    assert cli.main(args) == 0
    assert received == ([Path(args[0])] if args else [None])


def test_run_gui_schedules_initial_file(monkeypatch):
    scheduled = []
    opened = []
    app = SimpleNamespace(
        root=SimpleNamespace(after=lambda delay, callback: scheduled.append(callback)),
        check=opened.append,
        run=lambda: 0,
    )
    monkeypatch.setattr(gui, "_supports_tk", lambda: True)
    monkeypatch.setattr(gui, "App", lambda: app)
    assert gui.run_gui(Path("论文.docx")) == 0
    assert not opened
    scheduled[0]()
    assert opened == ["论文.docx"]


def test_fix_uses_displayed_report_after_source_changes(tmp_path, monkeypatch):
    path = tmp_path / "论文.docx"
    doc = Document()
    doc.add_heading("参考文献", level=1)
    original = "[1] 张三. 原始标题[J]. 学报, 2024"
    doc.add_paragraph(original)
    doc.save(path)
    app = gui.App.__new__(gui.App)
    callbacks = []
    app.root = SimpleNamespace(after=lambda delay, callback: callbacks.append(callback))
    app._job_id = 1
    app._last_path = str(path)
    app._entry_snapshot = ()

    # 只替换渲染控件；保留真实 _finish 的快照接收及过期任务判断。
    app._text_cache = ""
    app._md_cache = None
    widget = SimpleNamespace(
        configure=lambda **kw: None,
        pack=lambda **kw: None,
        delete=lambda *args: None,
        insert=lambda *args: None,
    )
    app.done_name = app.report = app.stats_bar = app.toolbar = widget
    app.btn_fix = app.btn_copy = app.btn_md = widget
    app.stat_values = dict.fromkeys(("total", "errors", "warns"), widget)
    app._show_drop_state = lambda mode: None
    app._set_stage = lambda *args, **kw: None
    app._set_status = lambda *args: None
    app._auto_fix_pending = False
    app._work(str(path), path.name, 1)
    callbacks.pop(0)()
    assert app._entry_snapshot == (original,)

    # 用户在 Word 中另存了新内容；当前报告仍应对应旧内容。
    doc.paragraphs[1].text = "[1] 李四. 后来修改的标题[J]. 学报, 2025."
    doc.save(path)
    finished = []
    app._finish_fix = lambda *args: finished.append(args)
    monkeypatch.setattr(
        gui.threading,
        "Thread",
        lambda target, args, daemon: SimpleNamespace(start=lambda: target(*args)),
    )
    app.generate_fix()
    callbacks.pop(0)()
    assert "原始标题" in finished[0][3]
    assert "后来修改" not in finished[0][3]


@pytest.mark.parametrize(
    "method,args",
    [
        ("_finish_error", (1, "旧错误")),
        ("_finish_fix_error", (1, "旧错误")),
        ("_finish_fix", (1, "旧修正", "旧统计", "旧内容", 0)),
    ],
)
def test_stale_callbacks_cannot_touch_new_job(method, args):
    app = gui.App.__new__(gui.App)
    app._job_id = 2
    app._fix_cache = "新内容"
    getattr(app, method)(*args)
    assert app._fix_cache == "新内容"


@pytest.mark.parametrize(
    "method,cache", [("export_md", "_md_cache"), ("export_fix_txt", "_fix_cache")]
)
def test_export_permission_error_has_recovery_message(monkeypatch, method, cache):
    app = gui.App.__new__(gui.App)
    setattr(app, cache, "待导出的内容")
    app._filedialog = SimpleNamespace(asksaveasfilename=lambda **kw: "report.txt")
    messages = []
    app._set_status = lambda message, color: messages.append(message)

    def denied(*args, **kwargs):
        raise PermissionError("文件被占用")

    monkeypatch.setattr(Path, "write_text", denied)
    getattr(app, method)()
    assert "导出失败" in messages[-1]
    assert "位置" in messages[-1]
    assert getattr(app, cache) == "待导出的内容"
