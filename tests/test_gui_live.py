"""显式启用的真实 Tk 集成检查：THESISLINT_RUN_GUI_TESTS=1 pytest tests/test_gui_live.py。"""

import os
from types import SimpleNamespace

import pytest
from docx import Document

from thesislint import gui

pytestmark = pytest.mark.skipif(
    os.environ.get("THESISLINT_RUN_GUI_TESTS") != "1",
    reason="真实 Tk 测试需桌面环境并显式启用",
)


@pytest.mark.parametrize("theme", ["dark", "light"])
def test_launch_analyze_fix_and_export(tmp_path, monkeypatch, theme):
    for key in ("THESISLINT_GUI_AUTODEMO", "THESISLINT_GUI_AUTOFIX", "THESISLINT_GUI_AUTOCLOSE_MS"):
        monkeypatch.delenv(key, raising=False)
    monkeypatch.setenv("THESISLINT_GUI_THEME", theme)
    path = tmp_path / "含空格的 论文.docx"
    output = tmp_path / "修正后.txt"
    doc = Document()
    doc.add_heading("参考文献", level=1)
    doc.add_paragraph("[1] 张三. 原始标题[J]. 学报, 2024\n[2] 李四. 文二[J]. 学报, 2023.")
    doc.save(path)
    real_app = gui.App
    errors = []
    finished = []

    def create_app():
        app = real_app()
        app.root.withdraw()
        assert hasattr(app.root, "drop_target_register")
        assert hasattr(app.root, "dnd_bind")
        app._filedialog = SimpleNamespace(asksaveasfilename=lambda **kw: str(output))

        def callback_failed(_kind, error, _traceback):
            errors.append(error)
            app.root.destroy()

        app.root.report_callback_exception = callback_failed

        def verify_fix():
            if not app._fix_cache:
                app.root.after(25, verify_fix)
                return
            assert "原始标题" in app._fix_cache
            assert "后来修改" not in app._fix_cache
            app.export_fix_txt()
            assert output.read_text(encoding="utf-8") == app._fix_cache
            app.toggle_theme()
            app.toggle_theme()
            assert app._theme_name == theme
            assert app.report.cget("background") == app._color("SURFACE")
            assert app.fix_text.tag_cget("warn", "foreground") == app._color("WARN")
            finished.append(True)
            app.root.destroy()

        def verify_report():
            if not app._entry_snapshot:
                app.root.after(25, verify_report)
                return
            assert len(app._entry_snapshot) == 2
            assert app._last_path == str(path)
            doc.paragraphs[1].text = "[1] 张三. 后来修改[J]. 学报, 2025."
            doc.save(path)
            app.generate_fix()
            app.root.after(25, verify_fix)

        def timeout():
            errors.append(TimeoutError("桌面检查和导出流程未在 10 秒内完成"))
            app.root.destroy()

        app.root.after(25, verify_report)
        app.root.after(10000, timeout)
        return app

    monkeypatch.setattr(gui, "App", create_app)
    assert gui.run_gui(path) == 0
    assert not errors, errors
    assert finished == [True]
