"""图形界面：双击 exe 或拖文件到图标时使用（Windows）。

深色「夜间自习室」风格，与网页版同一设计语言；v0.4.0 起与网页功能对齐：
拖拽进件、体检报告、一键生成修正后列表、复制/导出。

测试钩子（环境变量）：
- THESISLINT_GUI_AUTOCLOSE_MS  窗口 N 毫秒后自动关闭
- THESISLINT_GUI_AUTODEMO=1    启动后自动用示例论文跑一次体检
- THESISLINT_GUI_AUTOFIX=1     体检完成后自动生成修正列表
- THESISLINT_GUI_THEME=light   以浅色主题启动（默认 dark）
"""

from __future__ import annotations

import contextlib
import os
import sys
import threading
from pathlib import Path

STANDARD_URL = (
    "https://openstd.samr.gov.cn/bzgk/std/newGbInfo?"
    "hcno=C6CE52E55AC09B9C79A20AEA77CEDD14"
)

# ---- 配色：深色「夜间自习室」+ 浅色「白纸校样」----
DARK_THEME = {
    "BG": "#0b0e13",
    "SURFACE": "#0f141b",
    "RAISE": "#131a23",
    "LINE": "#1e2630",
    "LINE_HI": "#31405a",
    "TEXT": "#e6edf3",
    "MUTED": "#8b949e",
    "FAINT": "#56606b",
    "ACCENT": "#5b9bff",
    "ACCENT_HI": "#8ab8ff",
    "OK": "#7ee787",
    "WARN": "#e3b341",
    "ERR": "#ff7b72",
    "PAPER": "#182231",
    "ON_ACCENT": "#07111f",
    "OK_BG": "#1a3f2a",
    "OK_BG_HI": "#215237",
    "SELECT_BG": "#2d4a73",
}

LIGHT_THEME = {
    "BG": "#f6f8fa",
    "SURFACE": "#ffffff",
    "RAISE": "#eef2f7",
    "LINE": "#d8dee6",
    "LINE_HI": "#9fb0c5",
    "TEXT": "#1f2328",
    "MUTED": "#59636e",
    "FAINT": "#7d8996",
    "ACCENT": "#1f6feb",
    "ACCENT_HI": "#1158c7",
    "OK": "#1a7f37",
    "WARN": "#9a6700",
    "ERR": "#cf222e",
    "PAPER": "#dbeafe",
    "ON_ACCENT": "#f8fbff",
    "OK_BG": "#dafbe1",
    "OK_BG_HI": "#c7f0d2",
    "SELECT_BG": "#b6d7ff",
}

BG = DARK_THEME["BG"]
SURFACE = DARK_THEME["SURFACE"]
RAISE = DARK_THEME["RAISE"]
LINE = DARK_THEME["LINE"]
LINE_HI = DARK_THEME["LINE_HI"]
TEXT = DARK_THEME["TEXT"]
MUTED = DARK_THEME["MUTED"]
FAINT = DARK_THEME["FAINT"]
ACCENT = DARK_THEME["ACCENT"]
ACCENT_HI = DARK_THEME["ACCENT_HI"]
OK = DARK_THEME["OK"]
WARN = DARK_THEME["WARN"]
ERR = DARK_THEME["ERR"]
PAPER = DARK_THEME["PAPER"]
ON_ACCENT = DARK_THEME["ON_ACCENT"]
OK_BG = DARK_THEME["OK_BG"]
OK_BG_HI = DARK_THEME["OK_BG_HI"]
SELECT_BG = DARK_THEME["SELECT_BG"]


def _supports_tk() -> bool:
    try:
        import tkinter  # noqa: F401

        return True
    except Exception:  # noqa: BLE001  ImportError 或显示环境缺失
        return False


def _launched_from_explorer() -> bool:
    """是否由资源管理器启动（拖文件到图标 / 双击）。

    借助 kernel32.GetConsoleProcessList：用户手动开的控制台里还有 shell
    等其他进程；而资源管理器给控制台程序拉起的临时控制台里只有自己。
    没有控制台（Git Bash 管道/脚本调用）不可能是拖拽，必须视为命令行。
    """
    if os.name != "nt":
        return False
    try:
        import ctypes

        k32 = ctypes.windll.kernel32
        if not k32.GetConsoleWindow():
            return False
        buf = (ctypes.c_uint32 * 1)()
        return k32.GetConsoleProcessList(buf, 1) == 1
    except Exception:  # noqa: BLE001  刻意宽捕获：任一失败都走降级路径
        return False


def should_launch_gui(argv: list[str]) -> bool:
    """双模式分发决策。argv 永远是显式列表，不含本程序名。"""
    if not _supports_tk():
        return False
    if len(argv) == 0:
        return True  # 直接双击 exe
    if len(argv) == 1 and argv[0].lower().endswith(".docx") and not argv[0].startswith("-"):
        return _launched_from_explorer()
    return False


_DEMO_ENTRIES = [
    "[1] 刘知远. 大模型时代的自然语言处理[M]. 北京: 清华大学出版社, 2023: 15-20.",
    "[2] Vaswani A, Shazeer N, Parmar J, et al. Attention is all you need[C]//NeurIPS. 2017: 5998-6008.",
    "[3] 张三，李四. 一种深度学习方法[J]. 计算机学报，2024，47（3）：100-110",
    "[4] Brown T B, Mann B, Ryder N, Subbiah M, Kaplan J. Language models are few-shot learners[J]. NeurIPS, 2020.",
    "[5] 中国互联网络信息中心. 第53次中国互联网络发展状况统计报告[R/OL]. (2024-03-22)[2025-01-10]. https://www.cnnic.net.cn/report.html.",
]


def _result_meta(found: bool, errors: int, warns: int, total: int) -> tuple[str, str, str]:
    """返回结果条使用的（图标、说明、颜色）。"""
    if not found:
        return "—", "未找到参考文献章节", ERR
    if not total:
        return "—", "未提取到参考文献条目", WARN
    if errors:
        return "×", f"发现 {errors} 个错误", ERR
    if warns:
        return "!", f"发现 {warns} 个警告", WARN
    return "✓", "全部通过", OK


def _split_drop_paths(data: str, splitlist) -> list[str]:
    """按 Tcl 列表语法解析拖拽路径，兼容空格与多文件。"""
    try:
        return [str(item) for item in splitlist(data)]
    except Exception:  # noqa: BLE001  第三方 DnD 数据异常时保留单路径降级
        fallback = data.strip().strip("{}")
        return [fallback] if fallback else []


class App:
    def __init__(self) -> None:
        import tkinter as tk
        from tkinter import filedialog, messagebox, ttk
        from tkinter import font as tkfont

        self._tk = tk
        self._ttk = ttk
        self._filedialog = filedialog
        self._messagebox = messagebox
        self._theme_name = "dark"
        self._palette = DARK_THEME
        self._themed_buttons = []
        self._md_cache: str | None = None
        self._text_cache = ""
        self._report_views = {"all": "", "error": "", "warn": ""}
        self._report_filter = "all"
        self._fix_cache = ""
        self._last_path: str | None = None
        self._entry_snapshot: tuple[str, ...] = ()
        self._job_id = 0
        self._hero_mode = "idle"  # idle / busy / done
        self._pulse_step = False

        root = self._make_root(tk)
        self.root = root
        root.title("thesis-lint · 参考文献国标体检")
        root.geometry("940x700")
        root.minsize(780, 580)
        root.configure(bg=BG)

        families = set(tkfont.families())
        self._ui = "Microsoft YaHei UI" if "Microsoft YaHei UI" in families else "Microsoft YaHei"
        self._mono = "Cascadia Code" if "Cascadia Code" in families else "Consolas"

        self._style_ttk(ttk)
        self._build_header()
        self._build_steps()
        self._build_drop_zone()
        self._build_status()
        self._build_stats()
        self._build_toolbar()
        self._build_report_view()
        self._build_fix_panel()
        self._bind_shortcuts()
        self._set_stage(1)
        if os.environ.get("THESISLINT_GUI_THEME", "dark").lower() == "light":
            self._apply_theme("light")
        else:
            self._apply_titlebar_theme()
        self._pulse()

        # 测试钩子
        self._auto_fix_pending = os.environ.get("THESISLINT_GUI_AUTOFIX") == "1"
        if os.environ.get("THESISLINT_GUI_AUTODEMO") == "1":
            root.after(300, self._auto_demo)
        auto_close = os.environ.get("THESISLINT_GUI_AUTOCLOSE_MS")
        if auto_close and auto_close.isdigit():
            root.after(int(auto_close), root.destroy)

    # ---------- 窗口与主题 ----------

    def _make_root(self, tk):
        """优先 tkinterdnd2 的 Tk 获得整窗拖拽；失败退回普通 Tk。"""
        try:
            from tkinterdnd2 import DND_FILES, TkinterDnD

            root = TkinterDnD.Tk()
            root.drop_target_register(DND_FILES)
            root.dnd_bind("<<Drop>>", self._on_drop)
            return root
        except Exception:  # noqa: BLE001  tkinterdnd2 不可用时回退纯 Tk（拖拽降级为点击）
            return tk.Tk()

    def _style_ttk(self, ttk):
        style = ttk.Style(self.root)
        # 刻意宽捕获：主题不可用就用默认主题
        with contextlib.suppress(Exception):
            style.theme_use("clam")
        style.configure(
            "TScrollbar",
            background=self._color("RAISE"),
            troughcolor=self._color("BG"),
            bordercolor=self._color("BG"),
            arrowcolor=self._color("MUTED"),
            lightcolor=self._color("RAISE"),
            darkcolor=self._color("RAISE"),
        )
        style.map("TScrollbar", background=[("active", self._color("LINE_HI"))])

    def _color(self, token: str) -> str:
        return getattr(self, "_palette", DARK_THEME)[token]

    def _resolve_color(self, color: str) -> str:
        token = next((key for key, value in DARK_THEME.items() if value == color), None)
        return self._color(token) if token else color

    def _button_colors(self, kind: str) -> tuple[str, str, str, str]:
        tokens = {
            "accent": ("ACCENT", "ON_ACCENT", "ACCENT_HI", "ON_ACCENT"),
            "ghost": ("RAISE", "TEXT", "LINE_HI", "TEXT"),
            "warn": ("RAISE", "WARN", "LINE_HI", "WARN"),
            "ok": ("OK_BG", "OK", "OK_BG_HI", "OK"),
        }
        return tuple(self._color(token) for token in tokens[kind])

    def _paint_button(self, btn, kind: str, *, hovered: bool = False):
        bg, fg, hbg, hfg = self._button_colors(kind)
        btn.configure(
            bg=hbg if hovered else bg,
            fg=hfg if hovered else fg,
            activebackground=hbg,
            activeforeground=hfg,
            disabledforeground=self._color("FAINT"),
            highlightbackground=bg,
            highlightcolor=self._color("ACCENT_HI"),
        )

    def _hover(self, btn, normal, hovered):
        btn.bind("<Enter>", lambda _e: btn.configure(bg=hovered))
        btn.bind("<Leave>", lambda _e: btn.configure(bg=normal))

    def _button(self, parent, text, command, *, kind="ghost"):
        tk = self._tk
        bg, fg, hbg, hfg = self._button_colors(kind)
        btn = tk.Button(
            parent,
            text=text,
            command=command,
            font=(self._ui, 9),
            cursor="hand2",
            relief="flat",
            bd=0,
            padx=14,
            pady=5,
            bg=bg,
            fg=fg,
            activebackground=hbg,
            activeforeground=hfg,
            disabledforeground=FAINT,
            highlightthickness=1,
            highlightbackground=bg,
            highlightcolor=ACCENT_HI,
            takefocus=True,
            state="normal",
        )
        btn.bind(
            "<Enter>",
            lambda _e: (
                self._paint_button(btn, kind, hovered=True)
                if str(btn["state"]) == "normal"
                else None
            ),
        )
        btn.bind("<Leave>", lambda _e: self._paint_button(btn, kind))
        self._themed_buttons.append((btn, kind))
        return btn

    def _set_enabled(self, btn, enabled):
        btn.configure(state="normal" if enabled else "disabled")

    def _bind_shortcuts(self):
        self.root.bind_all("<Control-o>", lambda _e: self.pick_file())
        self.root.bind_all("<Control-O>", lambda _e: self.pick_file())
        self.root.bind_all("<Control-Shift-C>", lambda _e: self.copy_report())
        self.root.bind_all("<Control-t>", lambda _e: self.toggle_theme())
        self.root.bind_all("<Control-T>", lambda _e: self.toggle_theme())
        self.root.bind_all("<Escape>", lambda _e: self._collapse_fix_panel())

    def _walk_widgets(self, widget):
        yield widget
        for child in widget.winfo_children():
            yield from self._walk_widgets(child)

    def toggle_theme(self):
        self._apply_theme("light" if self._theme_name == "dark" else "dark")

    def _apply_theme(self, theme_name: str):
        old_palette = self._palette
        new_palette = LIGHT_THEME if theme_name == "light" else DARK_THEME
        old_colors = {value.lower(): key for key, value in old_palette.items()}
        color_options = (
            "background",
            "foreground",
            "activebackground",
            "activeforeground",
            "disabledforeground",
            "highlightbackground",
            "highlightcolor",
            "insertbackground",
            "selectbackground",
            "selectforeground",
        )

        for widget in self._walk_widgets(self.root):
            changes = {}
            for option in color_options:
                try:
                    current = str(widget.cget(option)).lower()
                except Exception:  # noqa: BLE001  不同 Tk 控件支持的颜色选项不同
                    continue
                token = old_colors.get(current)
                if token:
                    changes[option] = new_palette[token]
            if changes:
                widget.configure(**changes)

        self._theme_name = theme_name
        self._palette = new_palette
        for button, kind in self._themed_buttons:
            self._paint_button(button, kind)
        self._paint_report_filters()
        for text in (self.report, self.fix_text):
            text.tag_configure("err", foreground=self._color("ERR"))
            text.tag_configure("warn", foreground=self._color("WARN"))
            text.tag_configure("head", foreground=self._color("OK"))
            text.tag_configure("note", foreground=self._color("WARN"))
        self._style_ttk(self._ttk)
        self.btn_theme.configure(text="深色" if theme_name == "light" else "浅色")
        self._apply_titlebar_theme()

    def _apply_titlebar_theme(self):
        if os.name != "nt":
            return
        with contextlib.suppress(Exception):
            import ctypes

            value = ctypes.c_int(1 if self._theme_name == "dark" else 0)
            ctypes.windll.dwmapi.DwmSetWindowAttribute(
                self.root.winfo_id(), 20, ctypes.byref(value), ctypes.sizeof(value)
            )

    # ---------- 区块构建 ----------

    def _build_header(self):
        tk = self._tk
        from . import __version__

        head = tk.Frame(self.root, bg=BG)
        head.pack(fill="x", padx=26, pady=(18, 8))
        tk.Label(head, text="thesis", font=(self._ui, 12, "bold"), bg=BG, fg=TEXT).pack(side="left")
        tk.Label(head, text="·", font=(self._ui, 12, "bold"), bg=BG, fg=ACCENT).pack(side="left")
        tk.Label(head, text="lint", font=(self._ui, 12, "bold"), bg=BG, fg=TEXT).pack(side="left")
        tk.Label(
            head,
            text=f"  参考文献是否符合 GB/T 7714-2025 · v{__version__}",
            font=(self._ui, 9),
            bg=BG,
            fg=MUTED,
        ).pack(side="left", padx=(10, 0))
        tk.Label(head, text="本地解析 · 不上传", font=(self._ui, 8), bg=BG, fg=FAINT).pack(
            side="right"
        )
        self.btn_theme = self._button(head, "浅色", self.toggle_theme, kind="ghost")
        self.btn_theme.pack(side="right", padx=(0, 10))

    def _build_steps(self):
        tk = self._tk
        bar = tk.Frame(self.root, bg=BG)
        bar.pack(fill="x", padx=26, pady=(2, 8))
        self.step_labels = []
        for index, text in enumerate(("选择论文", "查看报告", "生成修正"), start=1):
            label = tk.Label(
                bar,
                text=f"{index}  {text}",
                font=(self._ui, 9, "bold"),
                bg=BG,
                fg=FAINT,
            )
            label.pack(side="left")
            self.step_labels.append(label)
            if index < 3:
                tk.Label(bar, text="  ─────  ", font=(self._mono, 8), bg=BG, fg=LINE_HI).pack(
                    side="left"
                )
        tk.Label(
            bar,
            text="Ctrl+O 选择文件 · Ctrl+T 切换明暗 · Esc 返回报告",
            font=(self._ui, 8),
            bg=BG,
            fg=FAINT,
        ).pack(side="right")

    def _set_stage(self, active: int, completed: int = 0):
        for index, label in enumerate(self.step_labels, start=1):
            color = (
                self._color("OK")
                if index <= completed
                else (self._color("ACCENT_HI") if index == active else self._color("FAINT"))
            )
            label.configure(fg=color)

    def _build_drop_zone(self):
        tk = self._tk
        self.drop = tk.Frame(
            self.root,
            bg=SURFACE,
            highlightthickness=1,
            highlightbackground=LINE,
            highlightcolor=LINE,
        )
        self.drop.pack(fill="x", padx=26, pady=(6, 4))

        # 海报态
        self.hero = tk.Frame(self.drop, bg=SURFACE)
        self.hero.pack(pady=34)
        tk.Label(
            self.hero,
            text="把论文拖进来，即刻体检",
            font=(self._ui, 16, "bold"),
            bg=SURFACE,
            fg=TEXT,
        ).pack()
        tk.Label(
            self.hero,
            text="拖入窗口、拖到 exe 图标、或点击选择 · 只支持 .docx（.doc 请先另存为 .docx）",
            font=(self._ui, 9),
            bg=SURFACE,
            fg=MUTED,
        ).pack(pady=(8, 16))
        row = tk.Frame(self.hero, bg=SURFACE)
        row.pack()
        self.btn_pick = self._button(row, "选择文件", self.pick_file, kind="accent")
        self.btn_pick.pack(side="left", padx=5)
        self.btn_demo = self._button(row, "用示例论文试一试", self.check_demo, kind="ghost")
        self.btn_demo.pack(side="left", padx=5)

        # 检查态：保持紧凑，让报告区尽早进入视野。
        self.busy = tk.Frame(self.drop, bg=SURFACE)
        self.busy_mark = tk.Label(
            self.busy,
            text="扫描中",
            font=(self._ui, 9, "bold"),
            bg=PAPER,
            fg=ACCENT_HI,
            padx=9,
            pady=3,
        )
        self.busy_mark.pack(side="left")
        self.busy_name = tk.Label(
            self.busy, text="", font=(self._mono, 10), bg=SURFACE, fg=TEXT, anchor="w"
        )
        self.busy_name.pack(side="left", padx=(12, 0))
        tk.Label(
            self.busy,
            text="正在定位参考文献并逐条检查…",
            font=(self._ui, 9),
            bg=SURFACE,
            fg=MUTED,
        ).pack(side="right")

        # 完成态（收拢成细条）
        self.done = tk.Frame(self.drop, bg=SURFACE)
        self.done_name = tk.Label(self.done, text="", font=(self._mono, 10), bg=SURFACE, fg=OK)
        self.done_name.pack(side="left")
        self.done_hint = tk.Label(
            self.done,
            text="点击或拖入新文件，重新检查",
            font=(self._ui, 9),
            bg=SURFACE,
            fg=ACCENT,
        )
        self.done_hint.pack(side="right")

        for w in (self.drop, self.hero, self.done, self.done_name, self.done_hint):
            w.bind("<Button-1>", lambda _e: self.pick_file())
        for child in self.hero.winfo_children():
            child.bind("<Button-1>", lambda _e: self.pick_file())

    def _build_status(self):
        tk = self._tk
        self.status = tk.Label(
            self.root, text="等待文件…", font=(self._ui, 9), bg=BG, fg=MUTED, anchor="w"
        )
        self.status.pack(fill="x", padx=28, pady=(6, 2))

    def _build_stats(self):
        tk = self._tk
        self.stats_bar = tk.Frame(self.root, bg=BG)
        self.stat_values = {}
        for key, label in (("total", "参考文献"), ("errors", "错误"), ("warns", "警告")):
            cell = tk.Frame(
                self.stats_bar, bg=RAISE, highlightthickness=1, highlightbackground=LINE
            )
            cell.pack(side="left", fill="x", expand=True, padx=(0, 8 if key != "warns" else 0))
            value = tk.Label(cell, text="0", font=(self._mono, 15, "bold"), bg=RAISE, fg=TEXT)
            value.pack(side="left", padx=(12, 7), pady=7)
            tk.Label(cell, text=label, font=(self._ui, 8), bg=RAISE, fg=MUTED).pack(side="left")
            self.stat_values[key] = value

    def _build_toolbar(self):
        outer = self._tk.Frame(self.root, bg=BG)
        self.toolbar = outer
        outer.pack(fill="x", padx=26, pady=(6, 8))
        bar = self._tk.Frame(outer, bg=BG)
        bar.pack(fill="x")
        self._tk.Label(bar, text="检查报告", font=(self._ui, 10, "bold"), bg=BG, fg=TEXT).pack(
            side="left", padx=(0, 14)
        )
        self.btn_copy = self._button(bar, "复制报告", self.copy_report, kind="ghost")
        self.btn_md = self._button(bar, "导出 Markdown", self.export_md, kind="ghost")
        self.btn_fix = self._button(bar, "生成修正后列表", self.generate_fix, kind="warn")
        for b in (self.btn_copy, self.btn_md, self.btn_fix):
            b.pack(side="left", padx=(0, 8))
            self._set_enabled(b, False)

        filter_bar = self._tk.Frame(outer, bg=BG)
        filter_bar.pack(fill="x", pady=(8, 0))
        self._tk.Label(
            filter_bar, text="报告索引", font=(self._ui, 8), bg=BG, fg=MUTED
        ).pack(side="left", padx=(0, 8))
        self.report_filter_buttons = {}
        for key, label in (("all", "全部"), ("error", "错误"), ("warn", "警告")):
            button = self._button(
                filter_bar,
                label,
                lambda selected=key: self._select_report_filter(selected),
                kind="ghost",
            )
            button.configure(padx=9, pady=3)
            button.pack(side="left", padx=(0, 5))
            self._set_enabled(button, False)
            self.report_filter_buttons[key] = button
        self.btn_basis = self._button(filter_bar, "规则依据 ↗", self.open_rule_basis, kind="ghost")
        self.btn_basis.configure(padx=9, pady=3)
        self.btn_basis.pack(side="right")

    def _make_code_view(self, parent, height):
        tk = self._tk
        frame = tk.Frame(parent, bg=LINE)
        text = tk.Text(
            frame,
            height=height,
            wrap="word",
            bg=SURFACE,
            fg=TEXT,
            insertbackground=TEXT,
            selectbackground="#2d4a73",
            selectforeground=TEXT,
            relief="flat",
            bd=0,
            highlightthickness=1,
            highlightbackground=LINE,
            font=(self._mono, 10),
            padx=16,
            pady=12,
            state="disabled",
        )
        text.pack(side="left", fill="both", expand=True)
        from tkinter import ttk

        sb = ttk.Scrollbar(frame, orient="vertical", command=text.yview)
        sb.pack(side="right", fill="y")
        text.configure(yscrollcommand=sb.set)
        text.tag_configure("err", foreground=ERR, font=(self._mono, 10, "bold"))
        text.tag_configure("warn", foreground=WARN, font=(self._mono, 10, "bold"))
        text.tag_configure("head", foreground=OK)
        text.tag_configure("note", foreground=WARN)
        return frame, text

    def _build_report_view(self):
        frame, text = self._make_code_view(self.root, 12)
        self.report_box, self.report = frame, text
        frame.pack(fill="both", expand=True, padx=26, pady=(0, 4))

    def _build_fix_panel(self):
        tk = self._tk
        self.fix_panel = tk.Frame(self.root, bg=BG)
        head = tk.Frame(self.fix_panel, bg=BG)
        head.pack(fill="x", pady=(10, 6))
        tk.Label(head, text="修正后的参考文献", font=(self._ui, 11, "bold"), bg=BG, fg=OK).pack(
            side="left"
        )
        self.fix_stats = tk.Label(head, text="", font=(self._ui, 9), bg=BG, fg=MUTED)
        self.fix_stats.pack(side="left", padx=(12, 0))

        frame, text = self._make_code_view(self.fix_panel, 9)
        self.fix_box, self.fix_text = frame, text
        frame.pack(fill="both", expand=True)

        tk.Label(
            self.fix_panel,
            text="⚠︎ 重新编号后请同步修改正文引用编号；粘贴回 Word 后建议再人工过一遍",
            font=(self._ui, 8),
            bg=BG,
            fg=FAINT,
            anchor="w",
        ).pack(fill="x", pady=(8, 4))
        row = tk.Frame(self.fix_panel, bg=BG)
        row.pack(fill="x")
        self.btn_fix_copy = self._button(row, "复制修正后条目", self.copy_fix, kind="ok")
        self.btn_fix_copy.pack(side="right")
        self.btn_fix_txt = self._button(row, "导出 txt", self.export_fix_txt, kind="ghost")
        self.btn_fix_txt.pack(side="right", padx=(0, 8))
        self.btn_fix_close = self._button(
            row, "返回检查报告", self._collapse_fix_panel, kind="ghost"
        )
        self.btn_fix_close.pack(side="right", padx=(0, 8))

    # ---------- 动效 ----------

    def _pulse(self):
        """拖拽区边框呼吸（海报态时）。"""
        try:
            if self._hero_mode == "idle":
                self._pulse_step = not self._pulse_step
                self.drop.configure(
                    highlightbackground=self._color("LINE_HI")
                    if self._pulse_step
                    else self._color("LINE")
                )
            elif self._hero_mode == "busy":
                self._pulse_step = not self._pulse_step
                self.drop.configure(
                    highlightbackground=self._color("ACCENT")
                    if self._pulse_step
                    else self._color("LINE_HI")
                )
            self.root.after(700, self._pulse)
        except self._tk.TclError:
            pass  # 窗口已关闭

    # ---------- 进件 ----------

    def pick_file(self):
        path = self._filedialog.askopenfilename(
            title="选择论文", filetypes=[("Word 文档", "*.docx"), ("所有文件", "*.*")]
        )
        if path:
            self.check(path)

    def _on_drop(self, event):
        paths = _split_drop_paths(event.data, self.root.tk.splitlist)
        if len(paths) != 1:
            self._messagebox.showwarning(
                "请一次拖入一个文件", "检测到多个文件。请只拖入一篇 Word 论文进行检查。"
            )
            return
        self.check(paths[0])

    def check_demo(self):
        import tempfile

        try:
            from docx import Document
        except Exception as exc:  # pragma: no cover  # noqa: BLE001
            self._set_status(f"无法生成示例：{exc}", ERR)
            return
        path = str(Path(tempfile.gettempdir()) / "thesislint_demo.docx")
        doc = Document()
        doc.add_heading("参考文献", level=1)
        for e in _DEMO_ENTRIES:
            doc.add_paragraph(e)
        doc.save(path)
        self.check(path, display_name="示例论文")

    def _auto_demo(self):
        self.check_demo()

    # ---------- 分析 ----------

    def _show_drop_state(self, mode: str):
        self._hero_mode = mode
        self.hero.pack_forget()
        self.busy.pack_forget()
        self.done.pack_forget()
        if mode == "idle":
            self.hero.pack(pady=34)
        elif mode == "busy":
            self.busy.pack(fill="x", padx=20, pady=13)
        else:
            self.done.pack(fill="x", padx=20, pady=13)

    def _clear_result_views(self):
        self._entry_snapshot = ()
        self._text_cache = ""
        self._report_views = {"all": "", "error": "", "warn": ""}
        self._report_filter = "all"
        self._md_cache = None
        self._fix_cache = ""
        self.stats_bar.pack_forget()
        self._show_report_panel()
        self.report.configure(state="normal")
        self.report.delete("1.0", "end")
        self.report.insert("end", "正在读取 Word 文档并定位参考文献…", ("note",))
        self.report.configure(state="disabled")
        self.btn_fix.configure(text="生成修正后列表")
        for button in getattr(self, "report_filter_buttons", {}).values():
            self._set_enabled(button, False)
        self._paint_report_filters()

    def _show_report_panel(self):
        self.fix_panel.pack_forget()
        self.report_box.pack_forget()
        self.report_box.pack(fill="both", expand=True, padx=26, pady=(0, 4), after=self.toolbar)

    def check(self, path: str, display_name: str | None = None):
        source = Path(path)
        if source.suffix.lower() != ".docx":
            self._messagebox.showwarning(
                "不支持的文件", "目前只支持 .docx 文件。\n如果是 .doc，请先用 Word 另存为 .docx。"
            )
            return
        if not source.is_file():
            self._set_status(f"文件不存在或无法读取：{source}", ERR)
            return

        self._job_id += 1
        job_id = self._job_id
        self._last_path = str(source)
        self._display_name = display_name or source.name
        self._clear_result_views()
        self._show_drop_state("busy")
        self.busy_name.configure(text=self._display_name)
        self._set_stage(2, completed=1)
        self._set_status(f"正在检查 {self._display_name} …", ACCENT)
        self.drop.configure(
            highlightbackground=self._color("LINE"), highlightcolor=self._color("LINE")
        )
        for b in (self.btn_copy, self.btn_md, self.btn_fix):
            self._set_enabled(b, False)
        self._set_enabled(self.btn_fix_copy, False)
        self._set_enabled(self.btn_fix_txt, False)
        threading.Thread(
            target=self._work,
            args=(str(source), self._display_name, job_id),
            daemon=True,
        ).start()

    def _work(self, path: str, display_name: str, job_id: int):
        try:
            from .cli import analyze
            from .report import build_markdown_report, build_text_report

            report = analyze(Path(path))
            text = build_text_report(report)
            text_views = {
                "all": text,
                "error": build_text_report(report, level="ERROR"),
                "warn": build_text_report(report, level="WARN"),
            }
            md = build_markdown_report(report)
            found = report.found_section
            if report.unavailable_reason:
                summary = report.unavailable_reason
                color = WARN if found else ERR
            else:
                summary = f"完成：{report.error_count} 个错误，{report.warn_count} 个警告"
                color = ERR if report.error_count else (WARN if report.warn_count else OK)
            stats = {
                "total": len(report.entries),
                "errors": report.error_count,
                "warns": report.warn_count,
            }
            entries = tuple(entry.text for entry in report.entries)
            self.root.after(
                0,
                lambda: self._finish(
                    job_id,
                    path,
                    display_name,
                    text,
                    md,
                    summary,
                    color,
                    found,
                    stats,
                    entries,
                    text_views,
                ),
            )
        except Exception as exc:  # 解析失败等  # noqa: BLE001  刻意宽捕获：任一失败都走降级路径
            detail = str(exc) or exc.__class__.__name__
            self.root.after(0, lambda: self._finish_error(job_id, detail))

    def _finish(
        self,
        job_id,
        path,
        display_name,
        text,
        md,
        summary,
        color,
        found,
        stats,
        entries,
        text_views,
    ):
        if job_id != self._job_id or self._last_path != path:
            return
        self._text_cache = text
        self._report_views = text_views
        self._report_filter = "all"
        self._md_cache = md
        self._entry_snapshot = entries

        icon, result_label, result_color = _result_meta(
            found, stats["errors"], stats["warns"], stats["total"]
        )
        self.done_name.configure(
            text=f"{icon}  {display_name} · {result_label}",
            fg=self._resolve_color(result_color),
        )
        self._show_drop_state("done")
        self._set_stage(2, completed=1)

        self._set_status(summary, color)
        if found:
            self.stat_values["total"].configure(text=str(stats["total"]), fg=self._color("TEXT"))
            self.stat_values["errors"].configure(
                text=str(stats["errors"]),
                fg=self._color("ERR") if stats["errors"] else self._color("MUTED"),
            )
            self.stat_values["warns"].configure(
                text=str(stats["warns"]),
                fg=self._color("WARN") if stats["warns"] else self._color("MUTED"),
            )
            self.stats_bar.pack(fill="x", padx=26, pady=(7, 2), before=self.toolbar)
        self._render_report_text(text)
        for button in getattr(self, "report_filter_buttons", {}).values():
            self._set_enabled(button, True)
        self._paint_report_filters()

        if found and stats["total"] > 0:
            self._set_enabled(self.btn_fix, True)
        self._set_enabled(self.btn_copy, True)
        self._set_enabled(self.btn_md, True)

        if self._auto_fix_pending and found and stats["total"] > 0:
            self._auto_fix_pending = False
            self.root.after(200, self.generate_fix)

    def _render_report_text(self, text: str):
        """把报告写入文本区，并按问题级别上色。"""
        self.report.configure(state="normal")
        self.report.delete("1.0", "end")
        for line in text.splitlines(True):
            if "[ERROR]" in line:
                tags = ("err",)
            elif "[WARN ]" in line or "[WARN]" in line:
                tags = ("warn",)
            elif line.startswith("体检结果"):
                tags = ("head",)
            else:
                tags = ()
            self.report.insert("end", line, tags)
        self.report.configure(state="disabled")

    def _select_report_filter(self, selected: str):
        text = getattr(self, "_report_views", {}).get(selected)
        if not text:
            return
        self._report_filter = selected
        self._render_report_text(text)
        self._paint_report_filters()

    def _paint_report_filters(self):
        for key, button in getattr(self, "report_filter_buttons", {}).items():
            self._paint_button(button, "ghost")
            if key == getattr(self, "_report_filter", "all"):
                button.configure(
                    bg=self._color("PAPER"),
                    fg=self._color("ACCENT_HI"),
                    highlightbackground=self._color("ACCENT"),
                )

    def open_rule_basis(self):
        import webbrowser

        if not webbrowser.open(STANDARD_URL):
            self._set_status("无法打开规则依据，请检查默认浏览器设置", WARN)

    def _finish_error(self, job_id: int, detail: str):
        if job_id != self._job_id:
            return
        self._last_path = None
        self._show_drop_state("idle")
        self._set_stage(1)
        self.report.configure(state="normal")
        self.report.delete("1.0", "end")
        self.report.insert(
            "end",
            f"无法读取这份文档。\n\n{detail}\n\n请确认文件未损坏且不是旧版 .doc 格式。",
            ("err",),
        )
        self.report.configure(state="disabled")
        self._set_status("检查失败，请选择另一份 .docx 文件", ERR)

    def _set_status(self, msg: str, color: str):
        self.status.configure(text=msg, fg=self._resolve_color(color))

    # ---------- 修正列表 ----------

    def generate_fix(self):
        if not self._entry_snapshot:
            return
        self._set_enabled(self.btn_fix, False)
        self._set_stage(3, completed=2)
        self._set_status("正在生成修正后列表 …", ACCENT)
        job_id = self._job_id
        threading.Thread(
            target=self._fix_work, args=(self._entry_snapshot, job_id), daemon=True
        ).start()

    def _fix_work(self, entries: tuple[str, ...], job_id: int):
        try:
            from .fixer import fix_entries

            result = fix_entries(list(entries))
            lines = result["lines"]
            s = result["summary"]
            header = (
                f"共 {s['total']} 条 · 自动修正 {s['auto_fixed']} 条 · "
                f"{s['manual_needed']} 条需人工处理" + (" · 已重新编号" if s["renumbered"] else "")
            )
            manual_lines = []
            for e in result["results"]:
                if e.unresolved:
                    manual_lines.append(
                        f"· 第 {e.index} 条（{e.original[:40]}…）：{'；'.join(e.unresolved)}"
                    )
            body = "\n".join(lines)
            fix_display = header + "\n\n" + body
            if manual_lines:
                fix_display += "\n\n以下条目无法全自动修正，请人工处理：\n" + "\n".join(
                    manual_lines
                )
            self.root.after(
                0,
                lambda: self._finish_fix(job_id, fix_display, header, body, len(manual_lines)),
            )
        except Exception as exc:  # noqa: BLE001  刻意宽捕获：任一失败都走降级路径
            detail = str(exc) or exc.__class__.__name__
            self.root.after(0, lambda: self._finish_fix_error(job_id, detail))

    def _finish_fix(self, job_id: int, fix_display: str, header: str, body: str, n_manual: int):
        if job_id != self._job_id:
            return
        self._fix_cache = body
        self.fix_stats.configure(text=header)
        self.fix_text.configure(state="normal")
        self.fix_text.delete("1.0", "end")
        for line in fix_display.splitlines(True):
            if line.startswith("共 "):
                tags = ("head",)
            elif line.startswith("· 第 "):
                tags = ("note",)
            else:
                tags = ()
            self.fix_text.insert("end", line, tags)
        self.fix_text.configure(state="disabled")
        self.report_box.pack_forget()
        self.fix_panel.pack(fill="both", expand=True, padx=26, pady=(0, 6))
        self._set_enabled(self.btn_fix_copy, True)
        self._set_enabled(self.btn_fix_txt, True)
        self.btn_fix.configure(text="重新生成修正后列表")
        self._set_enabled(self.btn_fix, True)
        self._set_stage(3, completed=3)
        self._set_status("修正后列表已生成，可一键复制", OK if n_manual == 0 else WARN)

    def _finish_fix_error(self, job_id: int, detail: str):
        if job_id != self._job_id:
            return
        self._set_stage(2, completed=1)
        self._set_status(f"生成失败：{detail}", ERR)
        self._set_enabled(self.btn_fix, True)

    def _collapse_fix_panel(self):
        if self.fix_panel.winfo_ismapped():
            self._show_report_panel()
            self._set_stage(2, completed=1)
            self.report.focus_set()

    # ---------- 导出 ----------

    def _save_export(self, target: str, content: str):
        try:
            Path(target).write_text(content, encoding="utf-8")
        except OSError as exc:
            self._set_status(f"导出失败，请选择其他保存位置或关闭占用文件的程序：{exc}", ERR)
            return
        self._set_status(f"已导出：{target}", OK)

    def copy_report(self):
        content = self._text_cache
        if not content:
            return
        self.root.clipboard_clear()
        self.root.clipboard_append(content)
        self._set_status("报告已复制到剪贴板", OK)

    def copy_fix(self):
        if not self._fix_cache:
            return
        self.root.clipboard_clear()
        self.root.clipboard_append(self._fix_cache)
        self._set_status("修正后列表已复制，去 Word 里粘贴吧", OK)

    def export_md(self):
        if not self._md_cache:
            return
        target = self._filedialog.asksaveasfilename(
            title="导出体检报告",
            defaultextension=".md",
            initialfile="参考文献体检报告.md",
            filetypes=[("Markdown", "*.md")],
        )
        if not target:
            return
        self._save_export(target, self._md_cache)

    def export_fix_txt(self):
        if not self._fix_cache:
            return
        target = self._filedialog.asksaveasfilename(
            title="导出修正后列表",
            defaultextension=".txt",
            initialfile="参考文献_修正后.txt",
            filetypes=[("文本文件", "*.txt")],
        )
        if not target:
            return
        self._save_export(target, self._fix_cache)

    def run(self) -> int:
        self.root.mainloop()
        return 0


def run_gui(initial_path: Path | None = None) -> int:
    """入口：返回进程退出码。"""
    if not _supports_tk():
        print(
            "当前环境缺少图形界面支持（tkinter），请改用命令行方式：thesislint 论文.docx",
            file=sys.stderr,
        )
        return 2
    if os.name == "nt":
        # 刻意宽捕获：DPI 感知失败不影响功能
        with contextlib.suppress(Exception):
            import ctypes

            ctypes.windll.shcore.SetProcessDpiAwareness(1)
    try:
        app = App()
    except Exception as exc:  # noqa: BLE001  刻意宽捕获：任一失败都走降级路径
        print(f"无法打开图形界面：{exc}\n可改用命令行：thesislint 论文.docx", file=sys.stderr)
        return 2
    if initial_path is not None:
        app.root.after(0, lambda: app.check(str(initial_path)))
    return app.run()
